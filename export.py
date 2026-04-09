"""
Export utilities — convert markdown discovery brief to Word (.docx).
"""

import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

SAP_BLUE = RGBColor(0x00, 0x66, 0xA1)
SAP_TEAL = RGBColor(0x3D, 0xBF, 0xB8)


def _set_heading_color(paragraph, color: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = color


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0066A1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_inline(para, text: str):
    """Add text to a paragraph with **bold** and *italic* inline formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                para.add_run(part)


def md_to_docx(markdown_text: str) -> bytes:
    """Convert markdown text to a styled Word document. Returns raw bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        i += 1

        if not stripped:
            continue

        if stripped == "---":
            _add_horizontal_rule(doc)

        elif stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=4)
            _set_heading_color(p, SAP_BLUE)

        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _set_heading_color(p, SAP_BLUE)

        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            _set_heading_color(p, SAP_BLUE)

        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _set_heading_color(p, SAP_BLUE)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, stripped[2:])

        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, re.sub(r"^\d+\. ", "", stripped))

        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        else:
            p = doc.add_paragraph()
            _parse_inline(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
